"""Parse uploaded stock account files into one login/account line per row."""

from __future__ import annotations

import csv
import io
import re
from typing import Any

_ALLOWED_EXT = {".txt", ".csv", ".xlsx", ".xls", ".pdf", ".docx", ".doc"}


def _clean_lines(text: str) -> list[str]:
    lines: list[str] = []
    for raw in (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw.strip()
        if line:
            lines.append(line)
    return lines


def _from_plain_text(data: bytes) -> list[str]:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return _clean_lines(data.decode(encoding))
        except UnicodeDecodeError:
            continue
    return _clean_lines(data.decode("utf-8", errors="ignore"))


def _from_csv(data: bytes) -> list[str]:
    text = data.decode("utf-8-sig", errors="ignore")
    reader = csv.reader(io.StringIO(text))
    out: list[str] = []
    for row in reader:
        cells = [str(c).strip() for c in row if str(c).strip()]
        if not cells:
            continue
        # Prefer joining non-empty cells so email,password stays useful.
        out.append(":".join(cells) if len(cells) > 1 else cells[0])
    return out


def _from_xlsx(data: bytes) -> list[str]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover
        raise ValueError("Excel support is not installed on the server") from exc
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    out: list[str] = []
    try:
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if not cells:
                    continue
                out.append(":".join(cells) if len(cells) > 1 else cells[0])
    finally:
        wb.close()
    return out


def _from_docx(data: bytes) -> list[str]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ValueError("Word support is not installed on the server") from exc
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(":".join(cells) if len(cells) > 1 else cells[0])
    # Also split multi-line paragraphs into separate stock lines.
    out: list[str] = []
    for part in parts:
        out.extend(_clean_lines(part))
    return out


def _from_pdf(data: bytes) -> list[str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ValueError("PDF support is not installed on the server") from exc
    reader = PdfReader(io.BytesIO(data))
    chunks: list[str] = []
    for page in reader.pages:
        chunks.append(page.extract_text() or "")
    return _clean_lines("\n".join(chunks))


def _from_legacy_doc(data: bytes) -> list[str]:
    """Best-effort extract from old .doc (OLE). Prefer .docx when possible."""
    # Many .doc files embed readable UTF-16 / ASCII runs — pull printable lines.
    text = data.decode("latin-1", errors="ignore")
    # Drop long binary noise; keep lines that look like accounts / emails.
    candidates = re.findall(r"[^\x00-\x08\x0b\x0c\x0e-\x1f]{4,}", text)
    joined = "\n".join(candidates)
    lines = _clean_lines(joined)
    # Keep lines that look useful (email, colon-separated, http, code-like).
    useful = [
        line
        for line in lines
        if ("@" in line or ":" in line or "http" in line.lower() or re.search(r"\d", line))
        and len(line) < 500
    ]
    if useful:
        return useful
    raise ValueError("Could not read this .doc file — save as .docx or .txt and try again")


async def extract_stock_lines_from_upload(upload: Any) -> list[str]:
    if upload is None or not (getattr(upload, "filename", None) or "").strip():
        return []
    name = (upload.filename or "").strip()
    lower = name.lower()
    ext = ""
    if "." in lower:
        ext = "." + lower.rsplit(".", 1)[-1]
    if ext not in _ALLOWED_EXT:
        raise ValueError("Bulk import supports .txt, .csv, .xlsx, .pdf, .docx (and .doc/.xls when readable)")

    data = await upload.read()
    if not data:
        return []

    if ext in {".txt"}:
        return _from_plain_text(data)
    if ext == ".csv":
        return _from_csv(data)
    if ext == ".xlsx":
        return _from_xlsx(data)
    if ext == ".xls":
        # openpyxl does not support legacy .xls
        raise ValueError("Legacy .xls is not supported — save as .xlsx or .csv")
    if ext == ".docx":
        return _from_docx(data)
    if ext == ".doc":
        # python-docx needs OOXML; try heuristic, else ask for docx.
        if data[:2] == b"PK":
            return _from_docx(data)
        return _from_legacy_doc(data)
    if ext == ".pdf":
        return _from_pdf(data)
    return []


def merge_login_lines(*chunks: str | list[str]) -> list[str]:
    out: list[str] = []
    for chunk in chunks:
        if isinstance(chunk, list):
            out.extend(_clean_lines("\n".join(chunk)))
        else:
            out.extend(_clean_lines(chunk or ""))
    return out
